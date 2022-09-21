from .imageprocessing import box, srceen_image
from .pakages.pytesseract import Output, pytesseract
from pdf2image import convert_from_path
import shutil
import docx
from autocorrect import Speller
import aspose.words as aw
import os
from bs4 import BeautifulSoup
from .consts import path_dic_data, tessdata_dir_config, poppler_path
from PIL import Image
import logging
from .imageprocessing import one_option_of_pro
from .textprocessing import  clear_character, create_text_file

pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract"


# for docker
# pytesseract.tesseract_cmd = "/usr/bin/tesseract"


def get_text_from_files(fullpath, cwd, file, lang, test_flag, ex):
    """
    get the file and his language ,process the file and return the text information
    """
    if not test_flag:
        with open(fullpath.encode('utf-8'), "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    text = ""
    if ex == "doc":
        logging.error("get doc file")
        path = (doc_to_doxc(fullpath))
        text = '\n'.join(get_text_from_docx(path))
    elif ex == "docx":
        logging.error("get docx file")
        text = '\n'.join(get_text_from_docx(fullpath))
    elif ex == "pdf":
        logging.error("get docx file")
        text = get_text_from_pdf(fullpath, cwd, lang)
    else:
        logging.error("get image file")
        get_text_from_image(fullpath, cwd, lang)
        with open(cwd + path_dic_data + '/result_text10.txt', encoding='utf8') as f:
            Lines = f.readlines()
            count = 0
            obj_data = {}
            for line in Lines:
                    count += 1
                    obj_data['Line'+str(count)] = line.strip()
        text = str(obj_data).replace('-\n', '')


    # for file in os.listdir(cwd + path_dic_data):
    #     os.remove(cwd + path_dic_data + file)
    return {"text": text}


def get_text_from_docx(fullpath):
    """
    get path of docx file and return the text in it
    """
    doc = docx.Document(fullpath)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return fullText


def doc_to_doxc(file_path):
    """
    convet doc file to docx
    """
    doc = aw.Document(file_path)
    cwd = os.getcwd()
    doc.save(cwd + "/tests/data/Output.docx")
    return cwd + "/tests/data/Output.docx"


def get_text_from_pdf(fullpath, cwd, lang):
    """
     get the path to the pdf file and return the text in it
     """
    # for docker
    pages = convert_from_path(fullpath, 500)
    # pages = convert_from_path(fullpath, 500, poppler_path=poppler_path)
    text = ''
    image_counter = 1
    for page in pages:
        filename = "page_" + str(image_counter) + ".jpg"
        data_path = cwd + path_dic_data + filename
        page.save(data_path, 'JPEG')
        image_counter = image_counter + 1
    filelimit = image_counter - 1
    for i in range(1, filelimit + 1):
        filename = "page_" + str(i) + ".jpg"
        try:
            without_effect = (pytesseract.image_to_string(cwd + path_dic_data + filename, output_type=Output.DICT,
                                                          config=tessdata_dir_config, lang=lang))
            # preprocessing image first option
            img =Image.open(cwd+path_dic_data+filename)
            one_option_of_pro(img, cwd)
            details = pytesseract.image_to_string(cwd + path_dic_data + '/final_img1.png', output_type=Output.DICT,
                                                  config=tessdata_dir_config, lang=lang)
            if len(without_effect['text']) > len(details['text']):
                details = without_effect
            details = details['text'].replace('\n', '  ')
            fix_data=''
            for word in details.split(" "):
                if lang == 'eng':
                    spelling = Speller(lang='en')
                    word = spelling(word)
                    fix_data += word + " "
                else:
                    fix_data += word + " "
            text += "page:" + filename + "data:" + str(details)
            os.remove(cwd + path_dic_data + filename)
        except Exception as e:
            print(e)
    return text


def get_text_from_image(fullpath, cwd, lang):
    """
    get fullpath to image and language which response to the image and return the text included in it
    """
    try:
        img = Image.open(fullpath)
    except:
        return ValueError
    # preprocessing image first option
    one_option_of_pro(img, cwd)
    # preprocessing image second option
    path = srceen_image(fullpath)
    box(path, cwd)
    # without processing image
    details1 = (pytesseract.image_to_data(fullpath, output_type=Output.DICT,
                                          config=tessdata_dir_config, lang=lang))
    # with the second option of process
    details3 = pytesseract.image_to_data(cwd + path_dic_data + r"/thresh.png", output_type=Output.DICT,
                                         config=tessdata_dir_config, lang=lang)
    # with the first option of process
    details5 = pytesseract.image_to_data(cwd + path_dic_data + '/final_img1.png', output_type=Output.DICT,
                                         config=tessdata_dir_config, lang=lang)

    # clear text ocr by remove unwanted character
    bad_chars = (';', ':', '!', "*", '|', '-', '%', '+', '=', '@', "+", ">", "<")
    arrays = [clear_character(details1, bad_chars), clear_character(details3, bad_chars),
              clear_character(details5, bad_chars)]
    list_of_value = []
    for data in arrays:
        list_of_value.append(len(data["text_clear"]))
    print(list_of_value)
    max_value = max(list_of_value)
    max_index = list_of_value.index(max_value)
    details = arrays[max_index]
    # tesst create text file
    create_text_file(details1, cwd, lang, 1, fullpath)
    create_text_file(details3, cwd, lang, 3, cwd + path_dic_data + r"/thresh.png")
    create_text_file(details5, cwd, lang, 5, cwd + path_dic_data + '/final_img1.png')
    data = create_text_file(details, cwd, lang, 10, fullpath)
    return data
